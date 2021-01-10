from datetime import datetime

from docx import Document
from docx.shared import Inches

from flaskr.report.plotly_chart_services import PlotlyChartService


class ReportFileService:
    def __init__(self):
        self.plotly_chart_service = PlotlyChartService()
        super().__init__()

    def create_file(self, report_data):
        document = Document()

        id = report_data['id']
        document.add_heading('Visium Load Test Run Report - ' + str(id), 0)

        document.add_heading('Test Setup Details', 1)

        start_date_time = datetime.fromtimestamp(report_data['startDateMs'] / 1000.0)
        end_date_time = datetime.fromtimestamp(report_data['endDateMs'] / 1000.0)
        elapsed_date_time = end_date_time - start_date_time

        records = (
            ('Locations Used', 'External'),
            ('Test Started', start_date_time.strftime('%c')),
            ('Test Finished', end_date_time.strftime('%c')),
            ('Elapsed', str(elapsed_date_time))
        )
        table = document.add_table(rows=0, cols=2)
        for key, value in records:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value

        document.add_heading('Overall Results', level=1)
        virtual_users_by_interval = report_data['virtualUsersByInterval']
        hit_point = len(virtual_users_by_interval)

        labels_date = self.get_date_labels(report_data['startDateMs'], report_data['endDateMs'], hit_point)

        chart_image_name = self.plotly_chart_service.save_throuhput_chart_image(virtual_users_by_interval,
                                                                                report_data['requestsByInterval'],
                                                                                labels_date)
        document.add_picture('images/' + chart_image_name)

        document.add_page_break()

        document.save('report-files/TestRun-{}.docx'.format(id))

    @staticmethod
    def get_date_labels(start_date_time, end_date_time, hit_point):
        labels = []
        if start_date_time != -1:
            time_diff = start_date_time - end_date_time
            aggregation_interval = int(time_diff / hit_point)

            current = start_date_time
            for x in range(hit_point):
                labels.append(current)
                current += aggregation_interval

        return labels
